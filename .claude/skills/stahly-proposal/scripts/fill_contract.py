"""
Stahly Professional Services Agreement (contract) filler + stapler.

Companion to build.py. Where build.py renders a polished proposal PDF, this script
fills a Stahly contract template (typically `Contract - General - <Office> for
Docusign.docx`) and optionally staples a proposal PDF onto the back so a single
combined file is ready for DocuSign.

USAGE
    python fill_contract.py --content contract.json --proposal proposal.pdf
                            [--out <output_dir>] [--template <docx_path>]
                            [--skill-dir <path>] [--no-staple]

Content JSON shape (see references/contract_fill_spec.md):
    {
      "agreement_date": "2026-05-14",
      "project_slug": "Hilger_to_Roy",         # used for output filenames
      "project_description": "Roy Substation to Hilger Substation LiDAR corridor survey.",
      "services_text": "As described in ...",
      "compensation_text": "As described in ...",
      "client": {
        "legal_name": "Fergus Electric Cooperative, Inc.",
        "address": "84423 US Highway 87, Lewistown, MT 59457",
        "signer_name": "Melanie Foran, P.E.",
        "signer_title": "Engineer"
      },
      "consultant": {
        "office": "Great Falls",                  # looked up in offices.json
        "signatory_name": "Aaron Kensinger, P.E.",
        "signatory_title": "Great Falls Regional Manager",   # short form for contract
        "pm_contact": "Ryan Harbach, P.L.S.",
        "pm_email": "rharbach@seaeng.com"
      }
    }

Output (default <output_dir> = parent of --proposal):
    <output_dir>/<YYYYMMDD>_<project_slug>_Contract_for_DocuSign.docx
    <output_dir>/<YYYYMMDD>_<project_slug>_Contract_for_DocuSign.pdf
    <output_dir>/<YYYYMMDD>_<project_slug>_Contract_AND_Proposal_for_DocuSign.pdf  (stapled)
    <output_dir>/<YYYYMMDD>_<project_slug>_Contract_for_DocuSign_log.md
"""
from __future__ import annotations
import argparse
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

# Force UTF-8 on Windows so non-ASCII status glyphs (✗ ✓ ⚠ → — etc.) don't
# crash with UnicodeEncodeError on cp1252 consoles.
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except AttributeError:
        import codecs
        sys.stdout = codecs.getwriter("utf-8")(sys.stdout.buffer, "strict")
        sys.stderr = codecs.getwriter("utf-8")(sys.stderr.buffer, "strict")

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement


# ============================================================
# DOCX EDIT HELPERS — same patterns proven on Hilger-to-Roy contract iter28
# ============================================================

def replace_para_text(para, new_text):
    """Replace full paragraph text in run 0, preserving its formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    r0 = para.runs[0]
    bold = r0.bold
    italic = r0.italic
    font_name = r0.font.name
    font_size = r0.font.size
    for r in para.runs:
        r.text = ""
    r0.text = new_text
    r0.bold = bold
    r0.italic = italic
    if font_name:
        r0.font.name = font_name
    if font_size:
        r0.font.size = font_size


def set_two_column_tab(para, x_inches=3.5):
    """Force a single LEFT-aligned tab stop at the given offset.

    Why: the Stahly contract template's default tab stops are 0.5" intervals.
    When the consultant's content is long (e.g. a long title), the second
    column (CLIENT side) gets pushed to a later default tab stop and the
    columns no longer line up under the CONSULTANT:/CLIENT: header. An
    explicit single tab stop at 3.5" keeps the right column always starting
    at the same x-position."""
    ts = para.paragraph_format.tab_stops
    ts.clear_all()
    ts.add_tab_stop(Inches(x_inches), WD_TAB_ALIGNMENT.LEFT)


def insert_empty_paragraph_before(para):
    """Inject a true empty paragraph in the document XML right before `para`.

    Why: setting `paragraph_format.space_before` is brittle — Word collapses
    adjacent space-after / space-before to the max of the two, and the
    template's preceding paragraph may already have space_after that absorbs
    the requested space_before. An empty paragraph is bulletproof — it
    occupies a full visual line regardless of surrounding styles."""
    empty = OxmlElement("w:p")
    para._element.addprevious(empty)


# ============================================================
# OFFICE / SIGNATORY LOOKUP — reuse the build.py source of truth
# ============================================================

def load_offices(skill_dir: Path):
    data = json.loads((skill_dir / "references" / "offices.json").read_text(encoding="utf-8"))
    return data["offices"]


def consultant_office_address(offices, office_name):
    if office_name not in offices:
        raise ValueError(f"Office not in directory: {office_name}")
    off = offices[office_name]
    return f"{off['street']}, {off['city_state_zip']}"


# ============================================================
# FILL ENGINE
# ============================================================

def fill_docx(template_path: Path, out_path: Path, ctx: dict, offices: dict):
    """Fill the contract template and save to out_path. Returns edits log."""
    shutil.copy2(str(template_path), str(out_path))
    doc = Document(str(out_path))

    client      = ctx["client"]
    consultant  = ctx["consultant"]
    cons_office = offices[consultant["office"]]
    cons_addr   = f"{cons_office['street']}, {cons_office['city_state_zip']}"
    agreement_d = _human_date(ctx["agreement_date"])

    edits = []

    # --- Header sentence ---
    for i, p in enumerate(doc.paragraphs):
        if "THIS AGREEMENT, entered into on" in p.text:
            new = (
                f"THIS AGREEMENT, entered into on {agreement_d}, by "
                f"{client['legal_name']} (hereafter “CLIENT”), located at "
                f"{client['address']} and Stahly Engineering & Associates, Inc. "
                f"(hereafter “CONSULTANT”), located at: {cons_addr}, "
                f"is hereby described as follows:"
            )
            replace_para_text(p, new)
            edits.append(f"P{i:03d}: header sentence filled")
            break

    # --- Section A: Project Description ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("A.") and "Project Description" in p.text:
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                if not doc.paragraphs[j].text.strip():
                    replace_para_text(doc.paragraphs[j], ctx["project_description"])
                    edits.append(f"P{j:03d}: Section A (Project Description) filled")
                    break
            break

    # --- Section B: Services ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("B.") and "professional services" in p.text:
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                if not doc.paragraphs[j].text.strip():
                    replace_para_text(doc.paragraphs[j], ctx["services_text"])
                    edits.append(f"P{j:03d}: Section B (Services) filled")
                    break
            break

    # --- Section C: Compensation ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("C.") and "compensate" in p.text:
            for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                if not doc.paragraphs[j].text.strip():
                    replace_para_text(doc.paragraphs[j], ctx["compensation_text"])
                    edits.append(f"P{j:03d}: Section C (Compensation) filled")
                    break
            break

    # --- Signature block (CONSULTANT/CLIENT/By/Name/Title) + PM/Contact + Email ---
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("CONSULTANT:") and "CLIENT" in text:
            replace_para_text(p, "CONSULTANT:\tCLIENT:")
            set_two_column_tab(p)
            for r in p.runs:
                r.font.size = Pt(12)
            edits.append(f"P{i:03d}: sig-block header at 3.5\" tab")
        elif text.startswith("By:"):
            replace_para_text(p, "By:  \tBy:  ")
            set_two_column_tab(p)
            for r in p.runs:
                r.font.size = Pt(12)
            edits.append(f"P{i:03d}: sig-block 'By:' at 3.5\" tab")
        elif text.startswith("Name:"):
            replace_para_text(p, f"Name:  {consultant['signatory_name']}\tName:  {client['signer_name']}")
            set_two_column_tab(p)
            for r in p.runs:
                r.font.size = Pt(12)
            edits.append(f"P{i:03d}: sig-block Name line at 3.5\" tab")
        elif text.startswith("Title:"):
            replace_para_text(p, f"Title:  {consultant['signatory_title']}\tTitle:  {client['signer_title']}")
            set_two_column_tab(p)
            for r in p.runs:
                r.font.size = Pt(12)
            edits.append(f"P{i:03d}: sig-block Title line at 3.5\" tab")
        elif text.startswith("Project Manager/Contact:"):
            # PM block becomes its own visual group:
            #  - empty paragraph above (bulletproof spacer)
            #  - 10pt font (smaller than 12pt sig block)
            #  - PM/Contact and Email on SEPARATE lines (soft line break)
            insert_empty_paragraph_before(p)

            if p.runs:
                orig_font_name = p.runs[0].font.name
                orig_bold = p.runs[0].bold
            else:
                orig_font_name = None
                orig_bold = False
            for r in p.runs:
                r.text = ""
            if p.runs:
                r0 = p.runs[0]
            else:
                r0 = p.add_run()
            r0.text = f"Project Manager/Contact:  {consultant['pm_contact']}"
            r0.font.size = Pt(10)
            r0.bold = orig_bold
            if orig_font_name:
                r0.font.name = orig_font_name

            r1 = p.add_run()
            r1.add_break()
            r1.add_text(f"Email Address:  {consultant['pm_email']}")
            r1.font.size = Pt(10)
            r1.bold = orig_bold
            if orig_font_name:
                r1.font.name = orig_font_name

            p.paragraph_format.tab_stops.clear_all()
            edits.append(f"P{i:03d}: PM/Contact split to 2 lines @ 10pt + empty-para spacer above")

    doc.save(str(out_path))
    return edits


def _human_date(iso):
    try:
        d = datetime.strptime(iso, "%Y-%m-%d")
        return d.strftime("%B %#d, %Y") if sys.platform.startswith("win") else d.strftime("%B %-d, %Y")
    except (ValueError, TypeError):
        return iso


# ============================================================
# DOCX -> PDF (Word automation)
# ============================================================

def docx_to_pdf(docx_path: Path, pdf_path: Path):
    """Open the docx in headless Word and SaveAs as PDF.

    Word automation requires ABSOLUTE paths — relative paths fail with "Sorry,
    we couldn't find your file" because Word's working directory is not the
    Python process's CWD."""
    import win32com.client
    docx_abs = str(docx_path.resolve())
    pdf_abs = str(pdf_path.resolve())
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(docx_abs)
        doc.SaveAs(pdf_abs, FileFormat=17)  # wdFormatPDF
        doc.Close()
    finally:
        word.Quit()


def staple_pdfs(contract_pdf: Path, proposal_pdf: Path, out_pdf: Path):
    from pypdf import PdfReader, PdfWriter
    w = PdfWriter()
    for src in [contract_pdf, proposal_pdf]:
        for page in PdfReader(str(src)).pages:
            w.add_page(page)
    with out_pdf.open("wb") as f:
        w.write(f)


# ============================================================
# LOG FILE
# ============================================================

def write_log(log_path: Path, ctx, template_path, proposal_pdf, contract_docx,
              contract_pdf, stapled_pdf, edits):
    lines = [
        f"# {log_path.stem.replace('_log', '')} - Contract Production Log",
        "",
        f"**Built:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"**Skill:** stahly-proposal / fill_contract.py",
        f"**Template:** `{template_path}`",
        f"**Proposal stapled:** `{proposal_pdf}`",
        f"**Outputs:**",
        f"  - docx: `{contract_docx}`",
        f"  - pdf:  `{contract_pdf}`",
        f"  - stapled: `{stapled_pdf}`",
        "",
        "## Identification",
        f"- Client: {ctx['client']['legal_name']}",
        f"- Client address: {ctx['client']['address']}",
        f"- Client signer: {ctx['client']['signer_name']}, {ctx['client']['signer_title']}",
        f"- Consultant signer: {ctx['consultant']['signatory_name']} ({ctx['consultant']['signatory_title']})",
        f"- PM/Contact: {ctx['consultant']['pm_contact']} <{ctx['consultant']['pm_email']}>",
        f"- Agreement date: {ctx['agreement_date']}",
        "",
        "## Project description (Section A)",
        f"> {ctx['project_description']}",
        "",
        "## Edits applied",
    ]
    for e in edits:
        lines.append(f"- {e}")
    log_path.write_text("\n".join(lines), encoding="utf-8")


# ============================================================
# CLI
# ============================================================

def main():
    ap = argparse.ArgumentParser(description="Fill a Stahly contract template and staple a proposal PDF behind it.")
    ap.add_argument("--content", required=True, help="Path to contract content JSON")
    ap.add_argument("--proposal", required=True, help="Path to the proposal PDF that will be stapled behind the contract")
    ap.add_argument("--template", default=None, help="Path to contract docx template (default: Great Falls template from \\\\Stahly\\stahly standards)")
    ap.add_argument("--out", default=None, help="Output directory (default: parent of --proposal)")
    ap.add_argument("--skill-dir", default=None, help="Skill directory (default: script's grandparent)")
    ap.add_argument("--no-staple", action="store_true", help="Skip the final staple step (just produce the filled contract)")
    args = ap.parse_args()

    skill_dir = Path(args.skill_dir) if args.skill_dir else Path(__file__).resolve().parent.parent
    content_path = Path(args.content)
    proposal_pdf = Path(args.proposal)
    if not proposal_pdf.exists():
        sys.exit(f"ERROR: proposal PDF not found: {proposal_pdf}")

    ctx = json.loads(content_path.read_text(encoding="utf-8"))

    # Resolve template path
    if args.template:
        template_path = Path(args.template)
    else:
        # Default per office; for now hardcode Great Falls. Extend later if needed.
        template_path = Path(r"\\Stahly\stahly standards\12-Forms_Templates\12.1-Forms\Great Falls\Contract - General - Great Falls for Docusign.docx")
    if not template_path.exists():
        sys.exit(f"ERROR: contract template not found: {template_path}")

    # Resolve output dir (create if missing)
    out_dir = Path(args.out) if args.out else proposal_pdf.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    # Compute filenames
    date_stamp = ctx["agreement_date"].replace("-", "")
    slug = ctx.get("project_slug", "Stahly_Project")
    contract_docx = out_dir / f"{date_stamp}_{slug}_Contract_for_DocuSign.docx"
    contract_pdf  = out_dir / f"{date_stamp}_{slug}_Contract_for_DocuSign.pdf"
    stapled_pdf   = out_dir / f"{date_stamp}_{slug}_Contract_AND_Proposal_for_DocuSign.pdf"
    log_path      = out_dir / f"{date_stamp}_{slug}_Contract_for_DocuSign_log.md"

    # Load offices for consultant address lookup
    offices = load_offices(skill_dir)
    if ctx["consultant"]["office"] not in offices:
        sys.exit(f"ERROR: office '{ctx['consultant']['office']}' not in offices.json")

    # 1. Fill the docx
    edits = fill_docx(template_path, contract_docx, ctx, offices)
    print(f"Filled: {contract_docx}")
    for e in edits:
        print(f"  {e}")

    # 2. docx -> pdf
    docx_to_pdf(contract_docx, contract_pdf)
    print(f"PDF: {contract_pdf}")

    # 3. Optional staple
    if not args.no_staple:
        staple_pdfs(contract_pdf, proposal_pdf, stapled_pdf)
        print(f"Stapled: {stapled_pdf}")

    # 4. Log
    write_log(log_path, ctx, template_path, proposal_pdf, contract_docx, contract_pdf,
              stapled_pdf if not args.no_staple else "(skipped)", edits)
    print(f"Log: {log_path}")


if __name__ == "__main__":
    main()
